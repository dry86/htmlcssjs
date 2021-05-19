from flask import Flask,request,jsonify,redirect,url_for
from flask_cors import CORS
from sklearn.feature_extraction.text import TfidfVectorizer
import numpy as np
from numpyencoder import NumpyEncoder
import json
import os
from scipy.linalg import norm
from werkzeug.utils import secure_filename
from docx import Document
from tqdm import tqdm
import zipfile
from zipfile import ZipFile
from pathlib import Path
import shutil
import jieba
from unrar import rarfile
from aip import AipNlp
from huaweicloud_nlp.NlpfClient import NlpfClient
from huaweicloud_nlp.HWNlpClientToken import HWNlpClientToken
from bs4 import BeautifulSoup
from win32com import client as wc
import pythoncom
from train_model import whole_dir
from train_model import get_corpus
from train_model import train_model
from train_model import get_stop_words
from train_model import jieba_cut
from gensim.models.doc2vec import Doc2Vec


###
app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = r"./upload"
CORS(app, supports_credentials=True)
#Baidu-NLP-API-KEY
APP_ID = '24093252'
API_KEY = 'VDeeMjwKexaYWpUTFQRV4CVa'
SECRET_KEY = '6viE1h8aKrVqxnLUaUcfGFjvRC0yqCjz'
client = AipNlp(APP_ID, API_KEY, SECRET_KEY)
#Huawei-NLP-API-KEY
tokenClient = HWNlpClientToken("QUST1708060227",  # domain name，用户的domain name
                               "QUST1708060227",  # 用户名
                               "Wws4615586", # 密码
                               "cn-north-4", # region
                               "090d360e77800f3a2fafc01ec9e1cfe1") # project_id
nlpfClient = NlpfClient(tokenClient)


#读txt 转 str
def read_txt_as_str(file_path):
    # 判断路径文件存在
    if not os.path.isfile(file_path):
        raise TypeError(file_path + " does not exist")
    try:
        all_the_text = open(file_path,encoding='gbk').read()
    except:
        all_the_text = open(file_path,encoding='utf-8').read()
    return all_the_text
#读docx 转 txt文件
def read_docx_as_txt(file_path,txt_save_path=app.config['UPLOAD_FOLDER']):
    file_name=os.path.basename(file_path)
    if file_name.endswith('.docx') or file_name.endswith('.doc'):
        document = Document(file_path)
        rename = os.path.splitext(file_name)
        txt_path = os.path.join(txt_save_path, rename[0]+'.txt')
        f = open(txt_path,'w', encoding='utf-8')
        for paragraph in tqdm(document.paragraphs):
                f.write(paragraph.text.strip()+'\n')
        f.close()
        return txt_path

def Baidu_NLP_API_simnet(text_1,text_2,options='BOW'):

    result = client.simnet(text_1,text_2)
    return json.dumps(result,ensure_ascii=False,sort_keys=True, indent=4, separators=(',', ': '))

def Huawei_NLP_API_similarity(text_1,text_2,lang='zh'):
    # 根据初始化Client章节选择认证方式构造完成nlpfClient后调用
    response = nlpfClient.get_text_similarity(text_1, text_2,lang)
    # 结果为code和json结构体
    print(response.code)
    #print(json.dumps(response.res,ensure_ascii=False)

    return json.dumps(response.res,ensure_ascii=False)


#---------------------------------------------------------------------------------------
@app.route('/upload_zip_whole',methods=["POST"])
def upload_zip_whole():
    # save zip
    file_dir_zip_raw = request.files['files']
    file_dir_zip_raw_path = os.path.join(app.config['UPLOAD_FOLDER'], file_dir_zip_raw.filename)
    file_dir_zip_raw.save(file_dir_zip_raw_path)


    #新建以提交zip为名称的dir，来存放解压的文件 注意文件“结构”
    print(file_dir_zip_raw.filename[::-1].split('.', 1)[-1][::-1])
    file_dir_zip_dir = os.path.join(app.config['UPLOAD_FOLDER'], file_dir_zip_raw.filename[::-1].split('.', 1)[-1][::-1])
    os.makedirs(file_dir_zip_dir)
    path_temp = os.path.join(file_dir_zip_dir , 'temp')
    os.makedirs(path_temp)
    print(file_dir_zip_dir)  #file dir #

    # read zip
    file_dir_zip = zipfile.ZipFile(file_dir_zip_raw_path)
    file_dir_zip_list = file_dir_zip.namelist()
    #print(file_dir_zip_list)

    pythoncom.CoInitialize() 
    data = [ ]
    data_count = 0


    def read_word_as_txt(file_path,txt_save_path):

        def save_doc_to_docx(rawpath):  # doc转docx
            word = wc.Dispatch("Word.Application")
            # 打开文件
            print(os.path.abspath(rawpath))
            doc = word.Documents.Open(os.path.abspath(rawpath))
            
            # # 将文件名与后缀分割
            rename = os.path.splitext(rawpath)
            print(rename[0])
            print(os.path.abspath(rename[0]))
            # 将文件另存为.docx
            doc.SaveAs( os.path.abspath(rename[0]) + '.docx', 12)  # 12表示docx格式
            doc.Close()


        file_name=os.path.basename(file_path)
        if not file_name.endswith('.pages') :
            if not file_name.endswith('.docx') :
                save_doc_to_docx(file_path)
                rename = os.path.splitext(file_path)
                file_path = rename[0] + '.docx'
            # document = Document(file_path)
            # f = open(txt_save_path,'w', encoding='utf-8')
            # for paragraph in tqdm(document.paragraphs):
            #         f.write(paragraph.text.strip()+'\n')
            # f.close()
            
            f = open(txt_save_path,'w', encoding='utf-8')
            document=ZipFile(file_path)
            xml_content=document.read("word/document.xml")
            wordObj=BeautifulSoup(xml_content.decode("utf-8"),"xml")
            texts=wordObj.findAll("w:t")
            # print(texts)
            for text in texts:
                #print(text.text)
                f.write(text.text)
            f.close()




    for f in file_dir_zip_list:
        # f_txt=file_dir_zip.open(f).read("word/document.xml")   #中文乱码
        # print(f_txt)
        #print(f)
        f_path = Path(file_dir_zip.extract(f,path_temp))
        #print(f_path)
        #截取文件名
        
        f_real_name = f.split("_")[0]     #xxx
        print(f_real_name)
        data_item = {'id':data_count,'name':f_real_name}
        data_count += 1
        data.append(data_item)

        txt_save_path = os.path.join(file_dir_zip_dir,f_real_name+'.txt')
        #print(txt_save_path)
        read_word_as_txt(f_path,txt_save_path)
        

    # 内存中有 data-dict 

    # 将一个文件夹下的txt 写入语料库
    whole_dir(file_dir_zip_dir)
    # get corpus
    x_train = get_corpus(file_dir_zip_dir)

    #model
    model = os.path.join(file_dir_zip_dir,"a_model_doc2vec")
    #训练model
    train_model(x_train,model)


    corpus_txt = file_dir_zip_dir +"\\a_zh_corpus_seg.txt"

    def use_model(test_text):
        model_dm = Doc2Vec.load(model)  # 不能每次都load model
        
        inferred_vector_dm = model_dm.infer_vector(line)
        sims = model_dm.docvecs.most_similar([inferred_vector_dm], topn=11)

        for i,score_data in enumerate(sims):
            if i == 0:
                
                continue
            print( data[i]['id'] )
            id,score=score_data
            if round(score-0.355,4) > 0:
                print(id)
                print(round(score,5))

   
        return sims
    
    stop_words = get_stop_words("stop_list.txt")
    model_dm = Doc2Vec.load(model)  # 不能每次都load model
    data_count = 0    
    txt_score = [ ]
    f_name_temp = ''
    for f in os.listdir(file_dir_zip_dir):
        if f.endswith('.txt') and not f.startswith('a'):
            text_path = os.path.join(file_dir_zip_dir,f)
            f_name = f.split('.')[0]
            # print(f_name)
            # print(data[data_count]['name'])
            if f_name == data[data_count]['name']:
                with open(text_path, "r", encoding='utf-8') as f_reader:
                    for line in f_reader:

                        word_list = jieba_cut(line, stop_words)
                        inferred_vector_dm = model_dm.infer_vector(word_list)
                        sims = model_dm.docvecs.most_similar([inferred_vector_dm], topn=11)
                        point_08 = []
                        point_07 = []
                        for i,score_data in enumerate(sims):
                            if i == 0:
                                continue
                            
                            id,score=score_data
                            

                            if round(score-0.8,4) > 0:
                                # print(id)    # 此id从 0 开始
                                # print(round(score,5))
                                point_08.append(score_data)
                            if round(score-0.7,4) > 0 and round(score-0.8,4) < 0:
                                # print(id)    # 此id从 0 开始
                                # print(round(score,5))
                                point_07.append(score_data)

                        txt_score_item = {'id':data_count,'name':f_name,'rate-0.8':point_08,'rate-0.7':point_07}    
                        # print(txt_score)
                        txt_score.append(txt_score_item)
            else:
                f_name_temp = f_name            
            data_count += 1 
                            




    
    data_json=json.dumps(txt_score, indent=4, ensure_ascii=False,cls=NumpyEncoder)
    with open(file_dir_zip_dir+ '\\a_txt_score.json','w',encoding='utf-8') as f:
        f.write(data_json) 
    
    # shutil.rmtree(path_temp)  #删除文件夹

    return "OK"  #返回的是 相似度数据



#---------------------------------------------------------------------------------------
#---------------------------------------------------------------------------------------
#---------------------------------------------------------------------------------------
@app.route('/upload_zip',methods=["POST"])
def upload_zip():
    #保存zip文件
    zip_file = request.files['files']
    zip_path = os.path.join(app.config['UPLOAD_FOLDER'], zip_file.filename)
    zip_file.save(zip_path)
    #从本地读取zip文件
    if zip_file.filename.endswith('.rar'):
        zip_file = rarfile.RarFile(zip_path,mode='r') #rar
        file = 'rar'
    else: 
        zip_file = zipfile.ZipFile(zip_path)          #zip
        file = 'zip'
    #解压
    zip_list = zip_file.namelist()
    zip_zh_list = []
    print(zip_list)
    for f in zip_list:
        if file == 'rar':
            # try:
            #     f_rename = f.split('\\',1)[1]  #分割出文件名
            # except:
            #     pass
            # print(f_rename)
            f_path = Path(zip_file.extract(f,app.config['UPLOAD_FOLDER']))
            print(f_path)
            read_docx_as_txt(f_path)
            
        elif file == 'zip':
            f_zh = f.encode('cp437').decode('gbk')
            zip_zh_list.append(f_zh)
            f_path = Path(zip_file.extract(f))
            f_zh_path = os.path.join(app.config['UPLOAD_FOLDER'], f_zh)
            #extracted_path.rename(f.encode('cp437').decode('gbk'))
            print(f_zh_path)
            shutil.move(f_path, f_zh_path)#将中文乱码名转化为中文正常文件名
            read_docx_as_txt(f_zh_path)#将word文档转化为txt文本
            print(zip_zh_list)
            
    if file == 'rar':
        return jsonify(zip_list)
    elif file == 'zip':
        return jsonify(zip_zh_list)
       
#---------------------------------------------------------------------------------------
   

@app.route('/compare_two_doc',methods=["GET"])
def compare_two_doc():


    d1 = request.args.get('a',type=str,default=None)
    d2 = request.args.get('b',type=str,default=None)

    d1_docx_path=os.path.join(app.config['UPLOAD_FOLDER'], d1)
    d1_txt_path=read_docx_as_txt(d1_docx_path)
    d1_str=read_txt_as_str(d1_txt_path)

    d2_docx_path=os.path.join(app.config['UPLOAD_FOLDER'], d2)
    d2_txt_path=read_docx_as_txt(d2_docx_path)
    d2_str=read_txt_as_str(d2_txt_path)
    #print(d1_str)
    result = Baidu_NLP_API_simnet(d1_str, d2_str)
    print(result)
    return jsonify(result)


@app.route('/upload',methods=["POST"]) # 方法要与前端一致
def upload():
    
    #file = request.files['files']  # Flask中获取文件
    #file_name = 
    for f in request.files.getlist('files'):
        #保存文件
        # print(file)
        print(f.filename)
        # print(file.name)
        #filename=secure_filename(file.filename)   
        f.save(os.path.join(app.config['UPLOAD_FOLDER'], f.filename))

    return 'success'


@app.route('/compare_string')
def tfidf_similarity():
    s1 = request.args.get('a',type=str,default=None)
    s2 = request.args.get('b',type=str,default=None)
    s3 = request.args.get('c',type=str,default=None)
    #print(s3)
    if s3 == 'Baidu':
        
        result = Baidu_NLP_API_simnet(s1, s2) #json
        result_dict = json.loads(result)
        result_score = result_dict["score"]
        print(result_score)

        return jsonify(result_score)

    if s3 == 'Huawei':
        result = Huawei_NLP_API_similarity(s1,s2)
        print(result)
        result_dict = json.loads(result)
        result_score = result_dict["similarity"]
        print(result_score)

        return jsonify(result_score)

    if s3 == 'TF-IDF':    
        def add_space(s):
            return ' '.join(list(s))
        # 将字中间加入空格
        s1, s2 = add_space(s1), add_space(s2)
        # 转化为TF矩阵
        cv = TfidfVectorizer(tokenizer=lambda s: s.split())
        corpus = [s1, s2]  #语料库
        #print(corpus)
        vectors = cv.fit_transform(corpus).toarray()

        # 计算TF系数
        result=np.dot(vectors[0], vectors[1]) / (norm(vectors[0]) * norm(vectors[1]))
        return jsonify(result)
        #计算余弦值cosθ

if __name__ == '__main__':
    app.debug=True
    app.run()